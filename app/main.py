
from __future__ import annotations

from pathlib import Path
from datetime import datetime
from io import BytesIO
import textwrap
import json

import numpy as np
import pandas as pd
import plotly.express as px
import streamlit as st

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader
from reportlab.platypus import (
    SimpleDocTemplate, BaseDocTemplate, PageTemplate, Frame,
    Paragraph, Spacer, Table, TableStyle, Image as RLImage,
    PageBreak, KeepTogether
)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# =============================================================================
# CONFIGURACIÓN GENERAL
# =============================================================================

APP_DIR = Path(__file__).resolve().parent.parent
ASSETS_DIR = APP_DIR / "assets"
EXPORTS_DIR = APP_DIR / "exports"
ASSETS_DIR.mkdir(exist_ok=True)
EXPORTS_DIR.mkdir(exist_ok=True)

LOGO_PATH = ASSETS_DIR / "logo_irrisal.jpg"

COMPANY_DEFAULTS = {
    "empresa": "Irrisal Consulting Ltda.",
    "direccion": "San Martín 553, oficina 901, Concepción, Región del Biobío, Chile.",
    "celular": "+56 9 6796 0884",
    "correo": "irrisalconsulting@gmail.com",
}

st.set_page_config(
    page_title="Pruebas de Bombeo - Irrisal Consulting",
    layout="wide"
)

# =============================================================================
# UTILIDADES
# =============================================================================

def clean_numeric_series(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce")


def safe_text(value, default: str = "Dato no informado") -> str:
    if value is None:
        return default
    if isinstance(value, float) and pd.isna(value):
        return default
    value = str(value).strip()
    return value if value else default


def fmt(value, suffix: str = "", decimals: int = 2, empty: str = "No evaluable") -> str:
    if value is None:
        return empty
    try:
        if pd.isna(value):
            return empty
    except Exception:
        pass
    if isinstance(value, (int, float, np.number)):
        return f"{float(value):.{decimals}f}{suffix}"
    return f"{value}{suffix}"


def df_numeric(df: pd.DataFrame, cols: list[str]) -> pd.DataFrame:
    out = df.copy()
    for c in cols:
        if c in out.columns:
            out[c] = pd.to_numeric(out[c], errors="coerce")
    return out


def add_warning(warnings: list[str], condition: bool, message: str):
    if condition:
        warnings.append(message)


# =============================================================================
# CÁLCULOS
# =============================================================================

def calculate_flow_stats(df: pd.DataFrame) -> dict:
    if df.empty or "Caudal_L_s" not in df.columns:
        return {"q_mean": None, "q_min": None, "q_max": None, "q_std": None, "q_var_pct": None, "is_constant": None}

    q = pd.to_numeric(df["Caudal_L_s"], errors="coerce").dropna()
    if q.empty:
        return {"q_mean": None, "q_min": None, "q_max": None, "q_std": None, "q_var_pct": None, "is_constant": None}

    q_mean = float(q.mean())
    q_min = float(q.min())
    q_max = float(q.max())
    q_std = float(q.std(ddof=0)) if len(q) > 1 else 0.0
    q_var_pct = ((q_max - q_min) / q_mean) * 100 if q_mean > 0 else None

    # Criterio simple: diferencia max-min <= 10% del caudal medio.
    is_constant = q_var_pct is not None and q_var_pct <= 10

    return {
        "q_mean": q_mean,
        "q_min": q_min,
        "q_max": q_max,
        "q_std": q_std,
        "q_var_pct": q_var_pct,
        "is_constant": is_constant,
    }


def calculate_pumped_volume(df: pd.DataFrame) -> float | None:
    if df.empty or not {"Tiempo_min", "Caudal_L_s"}.issubset(df.columns):
        return None

    data = df_numeric(df, ["Tiempo_min", "Caudal_L_s"]).dropna(subset=["Tiempo_min", "Caudal_L_s"])
    data = data.sort_values("Tiempo_min")
    if len(data) < 2:
        return None

    total = 0.0
    times = data["Tiempo_min"].to_numpy()
    flows = data["Caudal_L_s"].to_numpy()

    for i in range(1, len(data)):
        dt = times[i] - times[i - 1]
        if dt < 0:
            continue
        q_avg = (flows[i] + flows[i - 1]) / 2
        total += q_avg * dt * 60 / 1000

    return float(total)


def evaluate_stabilization(df: pd.DataFrame) -> dict:
    result = {
        "evaluable": False,
        "slope_cm_h": None,
        "meets": False,
        "message": "No evaluable: faltan datos suficientes.",
    }

    if df.empty or not {"Tiempo_min", "Nivel_m"}.issubset(df.columns):
        return result

    data = df_numeric(df, ["Tiempo_min", "Nivel_m"]).dropna(subset=["Tiempo_min", "Nivel_m"])
    data = data.sort_values("Tiempo_min")
    if len(data) < 2:
        return result

    duration = float(data["Tiempo_min"].max() - data["Tiempo_min"].min())
    if duration < 180:
        result["message"] = "No evaluable: se requieren al menos 180 minutos de datos para evaluar tendencia de estabilización."
        return result

    tmax = data["Tiempo_min"].max()
    last = data[data["Tiempo_min"] >= tmax - 180].copy()
    if len(last) < 2:
        result["message"] = "No evaluable: no hay suficientes puntos en los últimos 180 minutos."
        return result

    x_h = last["Tiempo_min"].to_numpy() / 60.0
    y_m = last["Nivel_m"].to_numpy()
    slope_m_h = np.polyfit(x_h, y_m, 1)[0]
    slope_cm_h = float(slope_m_h * 100)

    result["evaluable"] = True
    result["slope_cm_h"] = slope_cm_h
    result["meets"] = slope_cm_h <= 2

    if result["meets"]:
        result["message"] = "Presenta estabilización o franca tendencia según criterio ≤ 2 cm/h en los últimos 180 minutos."
    else:
        result["message"] = "No presenta estabilización según criterio ≤ 2 cm/h en los últimos 180 minutos."

    return result


def calculate_recovery(recovery_df: pd.DataFrame, static_level: float, final_dynamic_level: float) -> pd.DataFrame:
    data = recovery_df.copy()
    if data.empty or not {"Tiempo_min", "Nivel_m"}.issubset(data.columns):
        data["Recuperacion_pct"] = np.nan
        return data

    data = df_numeric(data, ["Tiempo_min", "Nivel_m"])
    denom = final_dynamic_level - static_level
    if denom <= 0:
        data["Recuperacion_pct"] = np.nan
        return data

    data["Recuperacion_pct"] = ((final_dynamic_level - data["Nivel_m"]) / denom) * 100
    data["Recuperacion_pct"] = data["Recuperacion_pct"].clip(lower=0, upper=100)
    return data


def time_to_recovery(df: pd.DataFrame, target_pct: float) -> float | None:
    if df.empty or "Recuperacion_pct" not in df.columns:
        return None
    valid = df_numeric(df, ["Tiempo_min", "Recuperacion_pct"]).dropna(subset=["Tiempo_min", "Recuperacion_pct"])
    reached = valid[valid["Recuperacion_pct"] >= target_pct]
    if reached.empty:
        return None
    return float(reached["Tiempo_min"].iloc[0])


def build_calculations(pumping_df: pd.DataFrame, recovery_df: pd.DataFrame, static_level: float) -> tuple[dict, pd.DataFrame]:
    pumping = df_numeric(pumping_df, ["Tiempo_min", "Nivel_m", "Caudal_L_s"])
    pumping_valid = pumping.dropna(subset=["Tiempo_min", "Nivel_m"]).sort_values("Tiempo_min")

    duration_min = None
    final_dynamic = None
    drawdown = None

    if not pumping_valid.empty:
        duration_min = float(pumping_valid["Tiempo_min"].max() - pumping_valid["Tiempo_min"].min())
        final_dynamic = float(pumping_valid["Nivel_m"].iloc[-1])
        drawdown = final_dynamic - static_level if final_dynamic >= static_level else None

    flow_stats = calculate_flow_stats(pumping)
    specific_capacity = flow_stats["q_mean"] / drawdown if drawdown and drawdown > 0 and flow_stats["q_mean"] is not None else None
    volume_m3 = calculate_pumped_volume(pumping)
    stabilization = evaluate_stabilization(pumping)

    recovery_with_pct = calculate_recovery(recovery_df, static_level, final_dynamic) if final_dynamic is not None else recovery_df.copy()
    if "Recuperacion_pct" not in recovery_with_pct.columns:
        recovery_with_pct["Recuperacion_pct"] = np.nan

    rec_valid = pd.to_numeric(recovery_with_pct["Recuperacion_pct"], errors="coerce").dropna()
    recovery_max = float(rec_valid.max()) if not rec_valid.empty else None

    calculations = {
        **flow_stats,
        "duration_min": duration_min,
        "final_dynamic": final_dynamic,
        "drawdown": drawdown,
        "specific_capacity": specific_capacity,
        "volume_m3": volume_m3,
        "stabilization_evaluable": stabilization["evaluable"],
        "stabilization_meets": stabilization["meets"],
        "slope_cm_h": stabilization["slope_cm_h"],
        "stabilization_message": stabilization["message"],
        "recovery_max": recovery_max,
        "t75": time_to_recovery(recovery_with_pct, 75),
        "t90": time_to_recovery(recovery_with_pct, 90),
        "t100": time_to_recovery(recovery_with_pct, 100),
    }

    return calculations, recovery_with_pct


# =============================================================================
# GRÁFICOS
# =============================================================================

def make_line_chart_image(
    df: pd.DataFrame,
    x_col: str,
    y_col: str,
    title: str,
    x_label: str,
    y_label: str,
    invert_y: bool = False,
) -> BytesIO | None:
    if df is None or df.empty or not {x_col, y_col}.issubset(df.columns):
        return None

    data = df_numeric(df, [x_col, y_col]).dropna(subset=[x_col, y_col]).sort_values(x_col)
    if len(data) < 2:
        return None

    fig, ax = plt.subplots(figsize=(8.0, 4.1))
    ax.plot(data[x_col], data[y_col], marker="o", linewidth=1.5, markersize=3.4)
    ax.set_title(title, fontsize=11)
    ax.set_xlabel(x_label, fontsize=9)
    ax.set_ylabel(y_label, fontsize=9)
    ax.grid(True, alpha=0.32)

    if invert_y:
        ax.invert_yaxis()

    fig.tight_layout()
    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=180)
    plt.close(fig)
    buffer.seek(0)
    return buffer


def image_flowable(image_buffer: BytesIO | None, width_cm: float = 16.5, height_cm: float = 8.4):
    if image_buffer is None:
        return Paragraph("Gráfico no disponible: datos insuficientes.", get_styles()["Body"])
    image_buffer.seek(0)
    img = rl_image_preserve_aspect(image_buffer, max_width_cm=width_cm, max_height_cm=height_cm)
    if img is None:
        return Paragraph("Gráfico no disponible: error al procesar imagen.", get_styles()["Body"])
    return img



# =============================================================================
# IMÁGENES CON PROPORCIÓN CONSERVADA
# =============================================================================

def _image_source_from_uploaded(uploaded_file):
    if uploaded_file is None:
        return None
    try:
        return BytesIO(uploaded_file.getvalue())
    except Exception:
        return None


def _scaled_dimensions(img_width_px, img_height_px, max_width_cm: float, max_height_cm: float):
    """
    Calcula dimensiones para insertar imagen sin deformarla.
    """
    max_w = max_width_cm * cm
    max_h = max_height_cm * cm

    if img_width_px <= 0 or img_height_px <= 0:
        return max_w, max_h

    scale = min(max_w / img_width_px, max_h / img_height_px)
    return img_width_px * scale, img_height_px * scale


def rl_image_preserve_aspect(source, max_width_cm: float, max_height_cm: float):
    """
    Crea un Image flowable de ReportLab preservando proporción.
    Acepta Path, str, BytesIO o UploadedFile convertido a BytesIO.
    """
    if source is None:
        return None

    try:
        if isinstance(source, Path):
            source_for_reader = str(source)
            source_for_image = str(source)
        elif isinstance(source, str):
            source_for_reader = source
            source_for_image = source
        elif isinstance(source, BytesIO):
            data = source.getvalue()
            source_for_reader = BytesIO(data)
            source_for_image = BytesIO(data)
        else:
            return None

        reader = ImageReader(source_for_reader)
        img_w, img_h = reader.getSize()
        width, height = _scaled_dimensions(img_w, img_h, max_width_cm, max_height_cm)
        return RLImage(source_for_image, width=width, height=height)
    except Exception:
        return None


def draw_header_logo(canvas, x, y, max_w, max_h):
    """
    Dibuja el logo en canvas sin deformarlo.
    """
    if not LOGO_PATH.exists():
        return False
    try:
        canvas.drawImage(
            str(LOGO_PATH),
            x,
            y,
            width=max_w,
            height=max_h,
            preserveAspectRatio=True,
            anchor="w",
            mask="auto",
        )
        return True
    except Exception:
        return False


# =============================================================================
# PDF: ESTILOS Y COMPONENTES
# =============================================================================

def get_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="CoverTitle", parent=styles["Title"], fontName="Helvetica-Bold",
        fontSize=18, leading=22, alignment=TA_CENTER,
        textColor=colors.HexColor("#006b2e"), spaceAfter=14
    ))
    styles.add(ParagraphStyle(
        name="CoverSubtitle", parent=styles["Normal"], fontName="Helvetica",
        fontSize=11, leading=16.5, alignment=TA_CENTER, spaceAfter=6
    ))
    styles.add(ParagraphStyle(
        name="SectionTitle", parent=styles["Heading2"], fontName="Helvetica-Bold",
        fontSize=13, leading=19.5, textColor=colors.HexColor("#006b2e"),
        spaceBefore=12, spaceAfter=8, keepWithNext=1
    ))
    styles.add(ParagraphStyle(
        name="Body", parent=styles["Normal"], fontName="Helvetica",
        fontSize=11, leading=16.5, alignment=TA_JUSTIFY
    ))
    styles.add(ParagraphStyle(
        name="TableBody", parent=styles["Normal"], fontName="Helvetica",
        fontSize=8.2, leading=10.5, alignment=TA_LEFT
    ))
    styles.add(ParagraphStyle(
        name="TableSmall", parent=styles["Normal"], fontName="Helvetica",
        fontSize=7.0, leading=8.6, alignment=TA_LEFT
    ))
    styles.add(ParagraphStyle(
        name="Small", parent=styles["Normal"], fontName="Helvetica",
        fontSize=8.5, leading=12.75, alignment=TA_JUSTIFY
    ))
    styles.add(ParagraphStyle(
        name="FigureCaption", parent=styles["Normal"], fontName="Helvetica",
        fontSize=8.5, leading=12.75, alignment=TA_CENTER,
        textColor=colors.HexColor("#555555"), spaceAfter=7
    ))
    return styles


def keep_paragraph(text_value: str, style):
    """
    Mantiene un párrafo completo junto. Evita títulos con una sola línea de contenido
    al final de página y el resto en la página siguiente.
    """
    return KeepTogether([Paragraph(text_value, style)])


def report_header_footer(canvas, doc, company: dict):
    canvas.saveState()
    width, height = A4

    green = colors.HexColor("#006b2e")
    gray = colors.HexColor("#444444")

    # Header: logo pequeño + nombre empresa
    logo_drawn = draw_header_logo(
        canvas,
        x=1.45 * cm,
        y=height - 1.05 * cm,
        max_w=1.25 * cm,
        max_h=0.72 * cm,
    )

    text_x = 2.85 * cm if logo_drawn else 1.5 * cm

    canvas.setFont("Helvetica-Bold", 7.4)
    canvas.setFillColor(green)
    canvas.drawString(text_x, height - 0.72 * cm, safe_text(company.get("empresa"), ""))

    canvas.setStrokeColor(green)
    canvas.setLineWidth(0.65)
    canvas.line(1.4 * cm, height - 1.22 * cm, width - 1.4 * cm, height - 1.22 * cm)

    # Footer: contacto + número de página
    canvas.setStrokeColor(colors.HexColor("#bbbbbb"))
    canvas.setLineWidth(0.35)
    canvas.line(1.4 * cm, 1.02 * cm, width - 1.4 * cm, 1.02 * cm)

    canvas.setFont("Helvetica", 6.5)
    canvas.setFillColor(gray)
    footer = f"{safe_text(company.get('direccion'), '')} | {safe_text(company.get('celular'), '')} | {safe_text(company.get('correo'), '')}"
    canvas.drawString(1.5 * cm, 0.68 * cm, footer[:130])
    canvas.drawRightString(width - 1.5 * cm, 0.68 * cm, f"Página {doc.page}")

    canvas.restoreState()

def make_table(data, col_widths=None, header=False, font_size=8.2, first_col_bold=True):
    wrapped = []
    styles = get_styles()
    table_style = styles["TableBody"]

    for row in data:
        wrapped.append([Paragraph(safe_text(cell, ""), table_style) for cell in row])

    table = Table(wrapped, colWidths=col_widths, repeatRows=1 if header else 0)
    style = [
        ("GRID", (0, 0), (-1, -1), 0.28, colors.HexColor("#999999")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), font_size),
        ("LEADING", (0, 0), (-1, -1), font_size * 1.25),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    if first_col_bold:
        style += [
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#edf7ef")),
        ]
    if header:
        style += [
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9ead3")),
        ]
    table.setStyle(TableStyle(style))
    return table

def df_to_pdf_table(df: pd.DataFrame, max_rows: int = 55, font_size: float = 7.0):
    styles = get_styles()
    if df is None or df.empty:
        return Paragraph("Sin datos ingresados.", styles["Body"])

    show = df.copy().head(max_rows).fillna("")

    # Formato específico: recuperación con dos decimales.
    if "Recuperacion_pct" in show.columns:
        show["Recuperacion_pct"] = pd.to_numeric(show["Recuperacion_pct"], errors="coerce").map(
            lambda x: "" if pd.isna(x) else f"{x:.2f}"
        )

    # Formato numérico para evitar decimales excesivos.
    for col in show.columns:
        if col != "Recuperacion_pct":
            numeric = pd.to_numeric(show[col], errors="coerce")
            if numeric.notna().sum() > 0 and numeric.notna().sum() >= len(show) * 0.5:
                show[col] = numeric.map(lambda x: "" if pd.isna(x) else f"{x:.3f}".rstrip("0").rstrip("."))

    data = [list(show.columns)] + show.astype(str).values.tolist()

    ncols = max(1, len(data[0]))
    col_width = min(16.5 / ncols, 4.5) * cm

    wrapped_data = []
    for row in data:
        wrapped_data.append([Paragraph(safe_text(cell, ""), styles["TableSmall"]) for cell in row])

    table = Table(wrapped_data, colWidths=[col_width] * ncols, repeatRows=1)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#999999")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9ead3")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), font_size),
        ("LEADING", (0, 0), (-1, -1), font_size * 1.25),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 2.5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 2.5),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))
    return table

def uploaded_image_flowable(uploaded_file, width_cm: float = 15.5, height_cm: float = 8.5):
    """
    Inserta imágenes subidas por el usuario sin deformarlas.
    """
    src = _image_source_from_uploaded(uploaded_file)
    if src is None:
        return None
    return rl_image_preserve_aspect(src, width_cm, height_cm)

def make_simple_well_scheme(capture: dict, stratigraphy_df: pd.DataFrame | None = None):
    """
    Genera un esquema constructivo referencial con formato más profesional.
    No inventa cribas, bomba ni niveles si los datos no fueron ingresados.
    Integra estratigrafía cuando existe información suficiente.
    """
    total_depth = capture.get("profundidad_total")
    static_level = capture.get("nivel_estatico")
    pump_depth = capture.get("profundidad_bomba")
    screen_from = capture.get("criba_desde")
    screen_to = capture.get("criba_hasta")

    try:
        total_depth = float(total_depth)
        if total_depth <= 0:
            return None
    except Exception:
        return None

    def _to_float_or_none(v):
        try:
            if v is None:
                return None
            if isinstance(v, str):
                v = v.strip().replace(",", ".")
                if v == "" or v.lower() in ["no informado", "dato no informado", "none", "nan"]:
                    return None
            out = float(v)
            if pd.isna(out) or out <= 0:
                return None
            return out
        except Exception:
            return None

    static_level = _to_float_or_none(static_level)
    pump_depth = _to_float_or_none(pump_depth)
    screen_from = _to_float_or_none(screen_from)
    screen_to = _to_float_or_none(screen_to)
    has_screen = screen_from is not None and screen_to is not None and screen_to > screen_from

    tipo = safe_text(capture.get("tipo"), "Captación")
    diam_ent = safe_text(capture.get("diametro_entubacion"), "")
    diam_perf = safe_text(capture.get("diametro_perforacion"), "")
    material = safe_text(capture.get("material_tuberia"), "")

    fig, ax = plt.subplots(figsize=(6.2, 8.7))
    ax.set_xlim(0, 14)
    ax.set_ylim(total_depth + max(2, total_depth * 0.08), -1.7)
    ax.axis("off")

    # Fondo general
    ax.add_patch(plt.Rectangle((0, -1.7), 14, total_depth + 4, facecolor="#fbfbf8", edgecolor="none"))

    # Escala de profundidad
    scale_x = 1.0
    ax.plot([scale_x, scale_x], [0, total_depth], color="#444444", linewidth=1.0)
    tick_step = 1 if total_depth <= 12 else 2 if total_depth <= 30 else 5
    ticks = np.arange(0, total_depth + 0.01, tick_step)
    for t in ticks:
        ax.plot([scale_x - 0.15, scale_x + 0.15], [t, t], color="#444444", linewidth=0.8)
        ax.text(scale_x - 0.25, t, f"{t:.0f}", ha="right", va="center", fontsize=7, color="#333333")
    ax.text(scale_x - 0.55, total_depth / 2, "Profundidad (m)", ha="center", va="center",
            fontsize=8, rotation=90, color="#333333")

    # Estratigrafía lateral
    strat_x0, strat_w = 1.55, 1.45
    has_valid_strat = False
    if stratigraphy_df is not None and not stratigraphy_df.empty:
        valid_strat = stratigraphy_df.copy()
        for col in ["Desde_m", "Hasta_m"]:
            if col in valid_strat.columns:
                valid_strat[col] = pd.to_numeric(valid_strat[col], errors="coerce")
        if {"Desde_m", "Hasta_m"}.issubset(valid_strat.columns):
            valid_strat = valid_strat.dropna(subset=["Desde_m", "Hasta_m"])
            valid_strat = valid_strat[valid_strat["Hasta_m"] > valid_strat["Desde_m"]].head(12)
            has_valid_strat = not valid_strat.empty
        else:
            valid_strat = pd.DataFrame()
    else:
        valid_strat = pd.DataFrame()

    if has_valid_strat:
        palette = ["#ead9c2", "#d8c8a8", "#c9d6b8", "#c8d8e4", "#d6d6d6", "#e2d2c4"]
        for i, (_, row) in enumerate(valid_strat.iterrows()):
            y0 = max(0, float(row["Desde_m"]))
            y1 = min(total_depth, float(row["Hasta_m"]))
            if y1 <= y0:
                continue
            color = palette[i % len(palette)]
            ax.add_patch(plt.Rectangle((strat_x0, y0), strat_w, y1 - y0,
                                       facecolor=color, edgecolor="#777777", linewidth=0.5))
            desc = safe_text(row.get("Descripcion", ""), "")
            label = f"{y0:g}-{y1:g} m"
            if desc:
                desc_short = desc[:24] + "…" if len(desc) > 24 else desc
                label += f"\n{desc_short}"
            ax.text(strat_x0 + strat_w/2, (y0+y1)/2, label, ha="center", va="center",
                    fontsize=6.3, color="#333333", wrap=True)
    else:
        ax.add_patch(plt.Rectangle((strat_x0, 0), strat_w, total_depth,
                                   facecolor="#f0e7d8", edgecolor="#777777", linewidth=0.5))
        ax.text(strat_x0 + strat_w/2, total_depth/2, "Estratigrafía\nno informada",
                ha="center", va="center", fontsize=7, color="#555555")
    ax.text(strat_x0 + strat_w/2, -0.65, "Estratigrafía", ha="center",
            va="bottom", fontsize=8, fontweight="bold", color="#006b2e")

    # Terreno y sello superficial
    ax.plot([0.7, 13.2], [0, 0], color="#5a3d22", linewidth=2.0)
    ax.add_patch(plt.Rectangle((3.2, -0.25), 4.1, 0.25, facecolor="#c9c9c9", edgecolor="#777777", linewidth=0.6))
    ax.text(9.2, -0.35, "Nivel de terreno", ha="left", va="center", fontsize=7.5, color="#333333")

    # Perforación y entubación
    bore_x0, bore_x1 = 4.2, 6.4
    casing_x0, casing_x1 = 4.72, 5.88
    center_x = (casing_x0 + casing_x1) / 2

    ax.add_patch(plt.Rectangle((bore_x0, 0), bore_x1-bore_x0, total_depth,
                               facecolor="#f6efe5", edgecolor="#7a7a7a", linewidth=1.0))
    ax.text((bore_x0+bore_x1)/2, -0.65, "Perforación", ha="center", va="bottom",
            fontsize=8, fontweight="bold", color="#006b2e")

    ax.add_patch(plt.Rectangle((casing_x0, 0), casing_x1-casing_x0, total_depth,
                               facecolor="#fdfdfd", edgecolor="#111111", linewidth=1.2))
    ax.add_patch(plt.Rectangle((casing_x0+0.12, 0), casing_x1-casing_x0-0.24, total_depth,
                               facecolor="#ffffff", edgecolor="#666666", linewidth=0.45))

    # Columna de agua y nivel estático
    water_start = static_level if static_level is not None and static_level < total_depth else total_depth
    if static_level is not None and water_start < total_depth:
        ax.add_patch(plt.Rectangle((casing_x0+0.18, water_start), casing_x1-casing_x0-0.36,
                                   total_depth-water_start, facecolor="#d8eef8",
                                   edgecolor="none", alpha=0.95))
        ax.plot([casing_x0 - 0.35, casing_x1 + 0.35], [static_level, static_level],
                color="#1877b7", linewidth=1.6)
        ax.text(7.05, static_level, f"Nivel estático: {static_level:.2f} m",
                ha="left", va="center", fontsize=7.5, color="#1877b7")
        ax.plot([casing_x1 + 0.05, 6.95], [static_level, static_level],
                color="#1877b7", linewidth=0.7)

    # Cribas / tramo ranurado
    if has_screen:
        sf, stt = max(0, screen_from), min(total_depth, screen_to)
        ax.add_patch(plt.Rectangle((casing_x0+0.12, sf), casing_x1-casing_x0-0.24, stt-sf,
                                   facecolor="#e8f4e8", edgecolor="#1f7a3a", linewidth=1.0))
        slot_count = max(5, int((stt - sf) / max(total_depth, 1) * 40))
        for y in np.linspace(sf + 0.12, stt - 0.12, slot_count):
            ax.plot([casing_x0+0.25, casing_x1-0.25], [y, y], color="#1f7a3a", linewidth=0.7)
        ax.text(7.05, (sf+stt)/2, f"Cribas / tramo ranurado:\n{sf:.1f} - {stt:.1f} m",
                ha="left", va="center", fontsize=7.3, color="#1f7a3a")
        ax.plot([casing_x1, 6.95], [(sf+stt)/2, (sf+stt)/2], color="#1f7a3a", linewidth=0.7)
    else:
        ax.text(7.05, total_depth*0.62, "Cribas / tramo filtrante:\nNo informado",
                ha="left", va="center", fontsize=7.3, color="#6a6a6a")
        ax.plot([casing_x1, 6.95], [total_depth*0.62, total_depth*0.62],
                color="#9a9a9a", linewidth=0.7, linestyle="--")

    # Bomba y columna
    if pump_depth is not None and pump_depth < total_depth:
        pump_h = max(total_depth * 0.035, 0.45)
        ax.add_patch(plt.Rectangle((center_x-0.25, pump_depth-pump_h/2), 0.5, pump_h,
                                   facecolor="#333333", edgecolor="#111111", linewidth=0.8))
        ax.add_patch(plt.Rectangle((center_x-0.08, 0), 0.16, max(0, pump_depth-pump_h/2),
                                   facecolor="#555555", edgecolor="none"))
        ax.text(7.05, pump_depth, f"Bomba: {pump_depth:.1f} m",
                ha="left", va="center", fontsize=7.5, color="#333333")
        ax.plot([center_x+0.25, 6.95], [pump_depth, pump_depth], color="#333333", linewidth=0.7)
    else:
        ax.text(7.05, total_depth*0.82, "Profundidad de bomba:\nNo informada",
                ha="left", va="center", fontsize=7.3, color="#6a6a6a")
        ax.plot([casing_x1, 6.95], [total_depth*0.82, total_depth*0.82],
                color="#9a9a9a", linewidth=0.7, linestyle="--")

    # Fondo del pozo
    ax.plot([casing_x0, casing_x1], [total_depth, total_depth], color="#111111", linewidth=1.2)
    ax.text(7.05, total_depth, f"Profundidad total: {total_depth:.1f} m",
            ha="left", va="center", fontsize=7.5, color="#333333")
    ax.plot([casing_x1, 6.95], [total_depth, total_depth], color="#333333", linewidth=0.7)

    # Caja de datos técnicos
    info_lines = [f"Tipo: {tipo}"]
    if diam_perf:
        info_lines.append(f"Diám. perforación: {diam_perf}")
    if diam_ent:
        info_lines.append(f"Diám. entubación: {diam_ent}")
    if material:
        info_lines.append(f"Material: {material}")
    info_text = "\n".join(info_lines)

    box_h = max(1.2, 0.45 + 0.45 * len(info_lines))
    ax.add_patch(plt.Rectangle((8.65, 0.55), 4.25, box_h,
                               facecolor="#ffffff", edgecolor="#b0b0b0", linewidth=0.7))
    ax.text(8.85, 0.88, info_text, ha="left", va="top", fontsize=7.4, color="#333333")

    # Título
    ax.text(7.0, -1.25, "Esquema constructivo referencial de la captación",
            ha="center", va="center", fontsize=10.5, fontweight="bold", color="#006b2e")
    ax.text(7.0, -0.92, "Representación no a escala horizontal; profundidad representada en metros.",
            ha="center", va="center", fontsize=7.2, color="#555555")

    fig.tight_layout(pad=0.4)
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=220, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf

# =============================================================================
# GENERACIÓN WORD
# =============================================================================

def add_docx_heading(document, text_value: str, level: int = 1):
    heading = document.add_heading(text_value, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = "Arial"
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 107, 46)
    heading.paragraph_format.line_spacing = 1.5
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.keep_together = True
    return heading


def add_docx_paragraph(document, text_value: str):
    p = document.add_paragraph()
    p.style = document.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    # Evita que un párrafo asociado a un título se parta dejando una sola línea en la página.
    p.paragraph_format.keep_together = True
    run = p.add_run(safe_text(text_value, ""))
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p



def set_docx_row_cant_split(row):
    """
    Evita, cuando Word lo respeta, que una fila de tabla se corte entre páginas.
    """
    try:
        trPr = row._tr.get_or_add_trPr()
        cant_split = row._tr._new_cantSplit()
        trPr.append(cant_split)
    except Exception:
        pass


def add_docx_table(document, rows: list[list], first_col_bold: bool = True):
    if not rows:
        return None

    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, row in enumerate(rows):
        try:
            set_docx_row_cant_split(table.rows[i])
        except Exception:
            pass
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = safe_text(value, "")
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    if first_col_bold and j == 0:
                        run.bold = True
    document.add_paragraph()
    return table

def add_df_docx_table(document, df: pd.DataFrame, max_rows: int = 80):
    if df is None or df.empty:
        add_docx_paragraph(document, "Sin datos ingresados.")
        return

    show = df.head(max_rows).fillna("").copy()

    if "Recuperacion_pct" in show.columns:
        show["Recuperacion_pct"] = pd.to_numeric(show["Recuperacion_pct"], errors="coerce").map(
            lambda x: "" if pd.isna(x) else f"{x:.2f}"
        )

    for col in show.columns:
        if col != "Recuperacion_pct":
            numeric = pd.to_numeric(show[col], errors="coerce")
            if numeric.notna().sum() > 0 and numeric.notna().sum() >= len(show) * 0.5:
                show[col] = numeric.map(lambda x: "" if pd.isna(x) else f"{x:.3f}".rstrip("0").rstrip("."))

    table = document.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr = table.rows[0].cells
    for j, col in enumerate(show.columns):
        hdr[j].text = str(col)
        for p in hdr[j].paragraphs:
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = "Arial"
                run.bold = True
                run.font.size = Pt(7)

    for _, row in show.iterrows():
        row_obj = table.add_row()
        try:
            set_docx_row_cant_split(row_obj)
        except Exception:
            pass
        cells = row_obj.cells
        for j, col in enumerate(show.columns):
            cells[j].text = str(row[col])
            for p in cells[j].paragraphs:
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7)

    document.add_paragraph()

def add_docx_picture_from_buffer(document, image_buffer: BytesIO | None, width_inches: float = 6.2):
    if image_buffer is None:
        add_docx_paragraph(document, "Imagen no disponible.")
        return

    try:
        image_buffer.seek(0)
        document.add_picture(image_buffer, width=Inches(width_inches))
    except Exception:
        add_docx_paragraph(document, "Imagen no disponible o no pudo insertarse.")


def add_docx_picture_from_upload(document, uploaded_file, width_inches: float = 6.2):
    if uploaded_file is None:
        return False
    try:
        document.add_picture(BytesIO(uploaded_file.getvalue()), width=Inches(width_inches))
        return True
    except Exception:
        return False



# =============================================================================
# TEXTOS NARRATIVOS AUTOMÁTICOS DEL INFORME
# =============================================================================

def _has_value(value) -> bool:
    if value is None:
        return False
    try:
        if pd.isna(value):
            return False
    except Exception:
        pass
    value = str(value).strip()
    return bool(value) and value.lower() not in ["dato no informado", "no informado", "none", "nan", "0", "0.0"]


def _phrase_value(value, suffix: str = "") -> str:
    if not _has_value(value):
        return "no informado"
    return f"{value}{suffix}"


def _format_duration_text(calculations: dict) -> str:
    duration = calculations.get("duration_min")
    if duration is None:
        return "una duración no informada"
    try:
        duration = float(duration)
    except Exception:
        return "una duración no informada"

    if duration >= 1440:
        return f"{duration:.0f} minutos, equivalentes a 24 horas"
    if duration >= 60:
        hours = duration / 60
        if abs(hours - round(hours)) < 0.01:
            return f"{duration:.0f} minutos, equivalentes a {hours:.0f} horas"
        return f"{duration:.0f} minutos, equivalentes a {hours:.1f} horas"
    return f"{duration:.0f} minutos"


def _format_location_text(project: dict, capture: dict) -> str:
    sector = safe_text(project.get("sector"), "")
    comuna = safe_text(project.get("comuna"), "")
    region = safe_text(project.get("region"), "")

    parts = []
    if sector:
        parts.append(f"sector {sector}")
    if comuna:
        parts.append(f"comuna de {comuna}")
    if region:
        parts.append(region)

    if parts:
        loc = ", ".join(parts)
    else:
        loc = "ubicación no informada"

    utm_este = safe_text(capture.get("utm_este"), "")
    utm_norte = safe_text(capture.get("utm_norte"), "")
    datum = safe_text(capture.get("datum"), "")
    huso = safe_text(capture.get("huso"), "")

    if utm_este and utm_norte:
        loc += f", con coordenadas UTM Este {utm_este} m y Norte {utm_norte} m"
        if datum or huso:
            loc += f", Datum/Huso {safe_text(datum, '')} {safe_text(huso, '')}".strip()
    return loc


def _stratigraphy_summary(stratigraphy_df: pd.DataFrame) -> str:
    if stratigraphy_df is None or stratigraphy_df.empty:
        return "No se ingresaron antecedentes estratigráficos detallados."

    rows = []
    for _, row in stratigraphy_df.head(5).iterrows():
        desde = row.get("Desde_m", "")
        hasta = row.get("Hasta_m", "")
        desc = row.get("Descripcion", "")
        if _has_value(desc):
            if _has_value(desde) and _has_value(hasta):
                rows.append(f"entre {desde} y {hasta} m se describe {str(desc).strip()}")
            else:
                rows.append(str(desc).strip())

    if not rows:
        return "No se ingresaron antecedentes estratigráficos detallados."

    return "De acuerdo con la estratigrafía ingresada, " + "; ".join(rows) + "."


def build_intro_text(project: dict, capture: dict, methodology: dict, calculations: dict) -> str:
    cliente = safe_text(project.get("cliente"), "el cliente")
    identificacion = safe_text(project.get("identificacion"), "la captación subterránea")
    location = _format_location_text(project, capture)

    return (
        f"El presente informe detalla los registros de mediciones efectuadas durante la prueba de bombeo "
        f"realizada en {identificacion}, correspondiente a {cliente}. La captación se ubica en {location}. "
        f"El documento reúne los antecedentes generales de la captación, la metodología aplicada, los registros "
        f"de nivel y caudal, la recuperación posterior al bombeo, los gráficos de comportamiento hidráulico y "
        f"las conclusiones técnicas derivadas de la información ingresada."
    )


def build_methodology_text(project: dict, capture: dict, equipment: dict, methodology: dict, calculations: dict) -> str:
    bomba = safe_text(equipment.get("bomba"), "equipo de bombeo no informado")
    potencia = safe_text(equipment.get("potencia"), "potencia no informada")
    profundidad_bomba = fmt(capture.get("profundidad_bomba"), " m")
    tuberia = safe_text(capture.get("tuberia_extraccion"), "tubería de extracción no informada")
    fecha = safe_text(methodology.get("fecha_prueba"), "fecha no informada")
    hora_inicio = safe_text(methodology.get("hora_inicio"), "hora no informada")
    hora_termino = safe_text(methodology.get("hora_termino"), "hora no informada")
    modo = safe_text(methodology.get("modo_prueba"), "modalidad no informada")
    caudal_obj = safe_text(methodology.get("caudal_objetivo"), "")
    caudal_prom = fmt(calculations.get("q_mean"), " L/s")
    duracion = _format_duration_text(calculations)
    medidor_caudal = safe_text(equipment.get("medidor_caudal") or methodology.get("metodo_caudal"), "instrumento de medición de caudal no informado")
    instrumento_nivel = safe_text(equipment.get("instrumento_nivel") or methodology.get("metodo_nivel"), "instrumento de medición de nivel no informado")
    frecuencia = safe_text(methodology.get("frecuencia"), "frecuencia no informada")

    caudal_text = f"a un caudal promedio calculado de {caudal_prom}"
    if _has_value(caudal_obj):
        caudal_text = f"considerando un caudal objetivo de {caudal_obj} y {caudal_text}"

    return (
        f"Para la ejecución de la prueba de bombeo se instaló {bomba}, con potencia {potencia}. "
        f"El equipo se dispuso a una profundidad de {profundidad_bomba}, utilizando {tuberia}. "
        f"El día {fecha}, a las {hora_inicio}, se dio inicio a la prueba bajo la modalidad {modo}, "
        f"{caudal_text}. La medición de caudal se efectuó mediante {medidor_caudal}, mientras que "
        f"los niveles dinámicos y de recuperación fueron controlados con {instrumento_nivel}. "
        f"Las mediciones se registraron con una frecuencia {frecuencia}, extendiéndose el bombeo por {duracion}. "
        f"Finalizado el bombeo a las {hora_termino}, se realizó el seguimiento de recuperación del nivel de agua "
        f"para evaluar la respuesta posterior de la captación."
    )


def build_capture_characteristics_text(capture: dict, stratigraphy_df: pd.DataFrame) -> str:
    tipo = safe_text(capture.get("tipo"), "captación subterránea")
    profundidad = fmt(capture.get("profundidad_total"), " m")
    diam_perf = safe_text(capture.get("diametro_perforacion"), "")
    diam_ent = safe_text(capture.get("diametro_entubacion"), "")
    material = safe_text(capture.get("material_tuberia"), "")
    nivel_estatico = fmt(capture.get("nivel_estatico"), " m")
    condicion = safe_text(capture.get("condicion"), "")
    criba_desde = capture.get("criba_desde")
    criba_hasta = capture.get("criba_hasta")
    tuberia_ciega = safe_text(capture.get("tuberia_ciega"), "")
    observaciones = safe_text(capture.get("observaciones"), "")

    desc = (
        f"La captación evaluada corresponde a {tipo}, con una profundidad total de {profundidad} "
        f"y nivel estático inicial de {nivel_estatico}."
    )

    detalles = []
    if _has_value(diam_perf):
        detalles.append(f"diámetro de perforación {diam_perf}")
    if _has_value(diam_ent):
        detalles.append(f"diámetro de entubación {diam_ent}")
    if _has_value(material):
        detalles.append(f"revestimiento o tubería de {material}")
    if _has_value(condicion):
        detalles.append(f"condición {condicion}")

    if detalles:
        desc += " La habilitación considera " + ", ".join(detalles) + "."

    if _has_value(criba_desde) and _has_value(criba_hasta):
        desc += f" El tramo ranurado o de cribas informado se extiende desde {criba_desde} m hasta {criba_hasta} m."
    else:
        desc += " No se cuenta con antecedentes informados de cribas, ranuras o tramo filtrante, por lo que este dato se declara como no informado."

    if _has_value(tuberia_ciega):
        desc += f" Se informa además tubería ciega o tramo sin ranurar: {tuberia_ciega}."

    desc += " " + _stratigraphy_summary(stratigraphy_df)

    if _has_value(observaciones):
        desc += f" Como observación de habilitación se registra: {observaciones}."

    return desc.strip()


def make_word_docx(
    company: dict,
    project: dict,
    capture: dict,
    stratigraphy_df: pd.DataFrame,
    equipment: dict,
    methodology: dict,
    pumping_df: pd.DataFrame,
    recovery_df: pd.DataFrame,
    calculations: dict,
    warnings: list[str],
    location_image=None,
    scheme_image=None,
    signature_image=None,
) -> bytes:
    """
    Genera informe Word editable (.docx) con la misma lógica del PDF.
    El Word permite corrección manual posterior antes de firmar o transformar a PDF.
    """
    document = Document()

    # Formato base Word: Arial 11, interlineado 1,5 en cuerpo.
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)


    # Márgenes
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    # Encabezado con logo pequeño
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO_PATH.exists():
        try:
            hp.add_run().add_picture(str(LOGO_PATH), width=Inches(0.55))
            hp.add_run("  ")
        except Exception:
            pass
    run = hp.add_run(safe_text(company.get("empresa"), ""))
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 107, 46)

    # Pie de página
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        f"{safe_text(company.get('direccion'), '')} | {safe_text(company.get('celular'), '')} | {safe_text(company.get('correo'), '')}"
    )
    fr.font.size = Pt(7)

    # Portada
    if LOGO_PATH.exists():
        try:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(LOGO_PATH), width=Inches(2.4))
        except Exception:
            pass

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("INFORME DE PRUEBA DE BOMBEO")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0, 107, 46)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(safe_text(project.get("identificacion"), "Captación subterránea"))
    sr.bold = True
    sr.font.size = Pt(12)

    add_docx_table(document, [
        ["Cliente / Beneficiario", project.get("cliente")],
        ["Proyecto", project.get("nombre_proyecto")],
        ["Sector / Predio", project.get("sector")],
        ["Comuna", project.get("comuna")],
        ["Región", project.get("region")],
        ["Fecha de prueba", methodology.get("fecha_prueba")],
        ["Fecha de emisión", datetime.now().strftime("%d-%m-%Y")],
        ["Consultor responsable", project.get("consultor")],
    ])

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(safe_text(company.get("empresa"))).bold = True
    add_docx_paragraph(document, safe_text(company.get("direccion")))
    add_docx_paragraph(document, f"Celular: {safe_text(company.get('celular'))} | Correo: {safe_text(company.get('correo'))}")

    document.add_page_break()

    add_docx_heading(document, "1. Introducción")
    add_docx_paragraph(document, build_intro_text(project, capture, methodology, calculations))

    add_docx_heading(document, "1.1 Antecedentes generales")
    add_docx_paragraph(document,
        "La información presentada a continuación corresponde a los antecedentes declarados para la prueba "
        "y a los datos técnicos registrados durante el ensayo. Estos antecedentes permiten contextualizar "
        "la ubicación de la captación, su habilitación y las condiciones bajo las cuales se efectuó la medición."
    )

    add_docx_heading(document, "1.2 Metodología de la prueba de bombeo")
    add_docx_paragraph(document, build_methodology_text(project, capture, equipment, methodology, calculations))

    add_docx_heading(document, "1.3 Características generales de la captación")
    add_docx_paragraph(document, build_capture_characteristics_text(capture, stratigraphy_df))

    add_docx_heading(document, "2. Síntesis de antecedentes generales")
    add_docx_table(document, [
        ["Cliente", project.get("cliente")],
        ["Proyecto", project.get("nombre_proyecto")],
        ["Identificación de captación", project.get("identificacion")],
        ["Sector / Predio", project.get("sector")],
        ["Comuna", project.get("comuna")],
        ["Región", project.get("region")],
        ["Consultor responsable", project.get("consultor")],
        ["Observaciones generales", project.get("observaciones")],
    ])

    add_docx_heading(document, "3. Ubicación y habilitación de la captación")
    cribas_text = (
        f"Desde {capture.get('criba_desde')} m hasta {capture.get('criba_hasta')} m"
        if capture.get("criba_desde") and capture.get("criba_hasta")
        else "No informado"
    )
    add_docx_table(document, [
        ["Tipo de captación", capture.get("tipo")],
        ["Coordenada UTM Norte", capture.get("utm_norte")],
        ["Coordenada UTM Este", capture.get("utm_este")],
        ["Datum / Huso", f"{safe_text(capture.get('datum'))} / {safe_text(capture.get('huso'))}"],
        ["Condición", capture.get("condicion")],
        ["Profundidad total", fmt(capture.get("profundidad_total"), " m")],
        ["Diámetro perforación", capture.get("diametro_perforacion")],
        ["Diámetro entubación", capture.get("diametro_entubacion")],
        ["Material / espesor tubería", capture.get("material_tuberia")],
        ["Altura sobre terreno", capture.get("altura_sobre_terreno")],
        ["Nivel estático inicial", fmt(capture.get("nivel_estatico"), " m")],
        ["Cribas / ranuras", cribas_text],
        ["Tubería ciega", capture.get("tuberia_ciega")],
        ["Profundidad de bomba", fmt(capture.get("profundidad_bomba"), " m")],
        ["Tubería extracción/succión", capture.get("tuberia_extraccion")],
    ])

    if location_image is not None:
        ok = add_docx_picture_from_upload(document, location_image, width_inches=6.2)
        if ok:
            add_docx_paragraph(document, "Figura 1. Croquis o imagen de ubicación de la captación.")

    add_docx_heading(document, "4. Esquema constructivo")
    if scheme_image is not None:
        ok = add_docx_picture_from_upload(document, scheme_image, width_inches=3.8)
        if ok:
            add_docx_paragraph(document, "Figura 2. Esquema constructivo de la captación.")
    else:
        scheme_buf = make_simple_well_scheme(capture, stratigraphy_df)
        add_docx_picture_from_buffer(document, scheme_buf, width_inches=3.6)
        add_docx_paragraph(document, "Figura 2. Esquema constructivo referencial de la captación.")

    add_docx_heading(document, "5. Estratigrafía")
    add_docx_paragraph(document, "La estratigrafía ingresada se presenta como antecedente descriptivo del material perforado o reconocido durante la habilitación.")
    add_df_docx_table(document, stratigraphy_df)

    add_docx_heading(document, "6. Equipos utilizados")
    add_docx_table(document, [
        ["Bomba", equipment.get("bomba")],
        ["Potencia", equipment.get("potencia")],
        ["Tubería de extracción", capture.get("tuberia_extraccion")],
        ["Medidor de caudal", equipment.get("medidor_caudal")],
        ["Instrumento de nivel", equipment.get("instrumento_nivel")],
        ["Generador", equipment.get("generador")],
    ])

    add_docx_heading(document, "7. Parámetros metodológicos registrados")
    add_docx_table(document, [
        ["Modo de prueba", methodology.get("modo_prueba")],
        ["Fecha de prueba", methodology.get("fecha_prueba")],
        ["Hora inicio bombeo", methodology.get("hora_inicio")],
        ["Hora término bombeo", methodology.get("hora_termino")],
        ["Duración registrada", fmt(calculations.get("duration_min"), " min", decimals=0)],
        ["Caudal objetivo", methodology.get("caudal_objetivo")],
        ["Método de medición de caudal", methodology.get("metodo_caudal")],
        ["Método de medición de niveles", methodology.get("metodo_nivel")],
        ["Frecuencia de medición", methodology.get("frecuencia")],
    ])

    add_docx_heading(document, "8. Resultados calculados")
    add_docx_table(document, [
        ["Caudal promedio", fmt(calculations.get("q_mean"), " L/s")],
        ["Caudal mínimo", fmt(calculations.get("q_min"), " L/s")],
        ["Caudal máximo", fmt(calculations.get("q_max"), " L/s")],
        ["Variación relativa de caudal", fmt(calculations.get("q_var_pct"), " %")],
        ["Nivel estático inicial", fmt(capture.get("nivel_estatico"), " m")],
        ["Nivel dinámico final", fmt(calculations.get("final_dynamic"), " m")],
        ["Abatimiento final", fmt(calculations.get("drawdown"), " m")],
        ["Caudal específico", fmt(calculations.get("specific_capacity"), " L/s/m")],
        ["Volumen bombeado", fmt(calculations.get("volume_m3"), " m³")],
        ["Pendiente final", fmt(calculations.get("slope_cm_h"), " cm/h")],
        ["Evaluación estabilización", calculations.get("stabilization_message")],
        ["Recuperación máxima", fmt(calculations.get("recovery_max"), " %")],
        ["Tiempo a 75% recuperación", fmt(calculations.get("t75"), " min", decimals=0)],
        ["Tiempo a 90% recuperación", fmt(calculations.get("t90"), " min", decimals=0)],
        ["Tiempo a 100% recuperación", fmt(calculations.get("t100"), " min", decimals=0)],
    ])

    add_docx_heading(document, "9. Gráficos")
    pump_chart = make_line_chart_image(pumping_df, "Tiempo_min", "Nivel_m", "Prueba de gasto constante", "Tiempo (min)", "Nivel/profundidad (m)", invert_y=True)
    add_docx_picture_from_buffer(document, pump_chart, width_inches=7.45)
    add_docx_paragraph(document, "Figura 3. Gráfico de prueba a caudal constante.")

    rec_chart = make_line_chart_image(recovery_df, "Tiempo_min", "Nivel_m", "Prueba de recuperación", "Tiempo (min)", "Nivel/profundidad (m)", invert_y=True)
    add_docx_picture_from_buffer(document, rec_chart, width_inches=6.6)
    add_docx_paragraph(document, "Figura 4. Gráfico de recuperación de nivel.")

    if recovery_df is not None and "Recuperacion_pct" in recovery_df.columns:
        rec_pct_chart = make_line_chart_image(recovery_df, "Tiempo_min", "Recuperacion_pct", "Porcentaje de recuperación", "Tiempo (min)", "Recuperación (%)", invert_y=False)
        add_docx_picture_from_buffer(document, rec_pct_chart, width_inches=6.6)
        add_docx_paragraph(document, "Figura 5. Porcentaje de recuperación acumulada.")

    add_docx_heading(document, "10. Tabla de prueba de gasto constante")
    add_df_docx_table(document, pumping_df)

    add_docx_heading(document, "11. Tabla de recuperación")
    add_df_docx_table(document, recovery_df)

    add_docx_heading(document, "12. Conclusiones")
    for c in generate_conclusions(capture, calculations, warnings, methodology.get("modo_prueba", "")):
        add_docx_paragraph(document, f"• {c}")

    # Bloque de firma profesional, centrado y separado del texto.
    spacer_p = document.add_paragraph()
    spacer_p.paragraph_format.space_before = Pt(28)
    spacer_p.paragraph_format.space_after = Pt(10)

    if signature_image is not None:
        try:
            sig_p = document.add_paragraph()
            sig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sig_p.paragraph_format.space_before = Pt(2)
            sig_p.paragraph_format.space_after = Pt(0)
            sig_p.add_run().add_picture(BytesIO(signature_image.getvalue()), width=Inches(2.35))
        except Exception:
            pass

    line_p = document.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_p.paragraph_format.space_before = Pt(0)
    line_p.paragraph_format.space_after = Pt(2)
    line_run = line_p.add_run("____________________________")
    line_run.font.name = "Arial"
    line_run.font.size = Pt(11)

    name_p = document.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after = Pt(0)
    name_run = name_p.add_run("David Gutiérrez Jara")
    name_run.font.name = "Arial"
    name_run.font.size = Pt(11)
    name_run.bold = True

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(0)
    title_run = title_p.add_run("Ingeniero Agrónomo")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(11)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


# =============================================================================
# GENERACIÓN PDF
# =============================================================================

def generate_conclusions(capture: dict, calculations: dict, warnings: list[str], test_mode: str) -> list[str]:
    tipo = safe_text(capture.get("tipo"), "")
    duration = calculations.get("duration_min")
    q_mean = calculations.get("q_mean")
    stab_meets = calculations.get("stabilization_meets")
    rec_max = calculations.get("recovery_max")

    conclusions = []

    if q_mean is not None:
        conclusions.append(
            f"Durante el periodo medido, la prueba se desarrolló con un caudal promedio de {fmt(q_mean, ' L/s')}."
        )

    if calculations.get("drawdown") is not None:
        conclusions.append(
            f"El abatimiento final calculado fue de {fmt(calculations.get('drawdown'), ' m')}, con un caudal específico de {fmt(calculations.get('specific_capacity'), ' L/s/m')}."
        )

    if calculations.get("stabilization_evaluable"):
        if stab_meets:
            conclusions.append(
                "La captación presenta estabilización o franca tendencia a estabilización bajo el criterio de variación ≤ 2 cm/h en los últimos 180 minutos evaluados."
            )
        else:
            conclusions.append(
                "La captación no presenta estabilización bajo el criterio de variación ≤ 2 cm/h en los últimos 180 minutos evaluados."
            )
    else:
        conclusions.append(
            "La estabilización no es evaluable con los datos disponibles."
        )

    if rec_max is not None:
        conclusions.append(
            f"La recuperación máxima observada alcanzó {fmt(rec_max, ' %')}. La interpretación de recuperación debe considerar la duración real del seguimiento posterior al bombeo."
        )

    if tipo == "Pozo profundo" and duration is not None and duration < 1440:
        conclusions.append(
            "La prueba corresponde a un ensayo de duración menor a 24 horas para pozo profundo. Sus resultados permiten una evaluación operativa del comportamiento de la captación durante el periodo medido, pero no reemplazan una prueba estándar de 24 horas cuando esta sea exigida."
        )

    if "abreviado" in test_mode.lower() or (duration is not None and duration <= 180):
        conclusions.append(
            "El ensayo abreviado debe interpretarse como antecedente técnico preliminar y sus conclusiones deben restringirse al periodo efectivamente medido."
        )

    conclusions.append(
        "El informe se basa exclusivamente en datos ingresados o importados por el usuario. El sistema no rellena mediciones faltantes ni presenta datos estimados como medidos."
    )

    return conclusions


def generate_recommendations(capture: dict, calculations: dict, warnings: list[str]) -> list[str]:
    recs = []
    tipo = safe_text(capture.get("tipo"), "")
    rec_max = calculations.get("recovery_max")

    if not calculations.get("stabilization_evaluable"):
        recs.append("Registrar al menos 180 minutos de mediciones continuas para evaluar tendencia de estabilización.")
    elif calculations.get("stabilization_meets"):
        recs.append("Mantener como referencia el caudal ensayado, siempre que las condiciones de operación y recuperación se mantengan similares.")
    else:
        recs.append("Evaluar una reducción del caudal de operación o repetir la prueba con mayor duración para definir un caudal más conservador.")

    if rec_max is None or rec_max < 75:
        recs.append("Extender la medición de recuperación hasta alcanzar al menos 75% de recuperación, idealmente hasta recuperación completa o estabilización clara.")

    if tipo == "Puntera":
        recs.append("Para punteras, habilitar piezómetro de control para medición de niveles, evitando interpretar niveles directamente desde la tubería de extracción.")

    if calculations.get("is_constant") is False:
        recs.append("Mejorar el control del caudal durante la prueba para asegurar condiciones de gasto constante.")

    recs.append("Conservar respaldo de planillas, fotografías, ubicación y equipos utilizados como anexos del expediente técnico.")

    return recs


def make_pdf(
    company: dict,
    project: dict,
    capture: dict,
    stratigraphy_df: pd.DataFrame,
    equipment: dict,
    methodology: dict,
    pumping_df: pd.DataFrame,
    recovery_df: pd.DataFrame,
    calculations: dict,
    warnings: list[str],
    location_image=None,
    scheme_image=None,
    signature_image=None,
) -> bytes:
    buffer = BytesIO()
    styles = get_styles()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.5 * cm,
        leftMargin=1.5 * cm,
        topMargin=1.85 * cm,
        bottomMargin=1.45 * cm,
    )

    def later_pages(canvas, doc_obj):
        report_header_footer(canvas, doc_obj, company)

    story = []

    # -------------------------------------------------------------------------
    # PORTADA
    # -------------------------------------------------------------------------
    if LOGO_PATH.exists():
        logo_flow = rl_image_preserve_aspect(LOGO_PATH, max_width_cm=6.2, max_height_cm=3.4)
        if logo_flow:
            story.append(logo_flow)
            story.append(Spacer(1, 0.35 * cm))

    story.append(Spacer(1, 0.25 * cm))
    story.append(Paragraph("INFORME DE PRUEBA DE BOMBEO", styles["CoverTitle"]))
    story.append(Paragraph(f"<b>{safe_text(project.get('identificacion'), 'Captación subterránea')}</b>", styles["CoverSubtitle"]))
    story.append(Spacer(1, 0.5 * cm))

    cover_data = [
        ["Cliente / Beneficiario", project.get("cliente")],
        ["Proyecto", project.get("nombre_proyecto")],
        ["Sector / Predio", project.get("sector")],
        ["Comuna", project.get("comuna")],
        ["Región", project.get("region")],
        ["Fecha de prueba", methodology.get("fecha_prueba")],
        ["Fecha de emisión", datetime.now().strftime("%d-%m-%Y")],
        ["Consultor responsable", project.get("consultor")],
    ]
    story.append(make_table(cover_data, col_widths=[5.2 * cm, 11.2 * cm]))
    story.append(Spacer(1, 1.0 * cm))
    story.append(Paragraph(f"<b>{safe_text(company.get('empresa'))}</b>", styles["CoverSubtitle"]))
    story.append(Paragraph(safe_text(company.get("direccion")), styles["CoverSubtitle"]))
    story.append(Paragraph(f"Celular: {safe_text(company.get('celular'))} | Correo: {safe_text(company.get('correo'))}", styles["CoverSubtitle"]))
    story.append(PageBreak())

    # -------------------------------------------------------------------------
    # 1. INTRODUCCIÓN Y TEXTOS NARRATIVOS
    # -------------------------------------------------------------------------
    story.append(Paragraph("1. Introducción", styles["SectionTitle"]))
    story.append(keep_paragraph(build_intro_text(project, capture, methodology, calculations), styles["Body"]))

    story.append(Paragraph("1.1 Antecedentes generales", styles["SectionTitle"]))
    story.append(Paragraph(
        "La información presentada a continuación corresponde a los antecedentes declarados para la prueba "
        "y a los datos técnicos registrados durante el ensayo. Estos antecedentes permiten contextualizar "
        "la ubicación de la captación, su habilitación y las condiciones bajo las cuales se efectuó la medición.",
        styles["Body"]
    ))

    story.append(Paragraph("1.2 Metodología de la prueba de bombeo", styles["SectionTitle"]))
    story.append(keep_paragraph(build_methodology_text(project, capture, equipment, methodology, calculations), styles["Body"]))

    story.append(Paragraph("1.3 Características generales de la captación", styles["SectionTitle"]))
    story.append(keep_paragraph(build_capture_characteristics_text(capture, stratigraphy_df), styles["Body"]))

    # -------------------------------------------------------------------------
    # 2. SÍNTESIS DE ANTECEDENTES GENERALES
    # -------------------------------------------------------------------------
    story.append(Paragraph("2. Síntesis de antecedentes generales", styles["SectionTitle"]))
    general_data = [
        ["Cliente", project.get("cliente")],
        ["Proyecto", project.get("nombre_proyecto")],
        ["Identificación de captación", project.get("identificacion")],
        ["Sector / Predio", project.get("sector")],
        ["Comuna", project.get("comuna")],
        ["Región", project.get("region")],
        ["Consultor responsable", project.get("consultor")],
        ["Observaciones generales", project.get("observaciones")],
    ]
    story.append(make_table(general_data, col_widths=[5.2 * cm, 11.2 * cm]))

    # -------------------------------------------------------------------------
    # 3. UBICACIÓN Y HABILITACIÓN
    # -------------------------------------------------------------------------
    story.append(Paragraph("3. Ubicación y habilitación de la captación", styles["SectionTitle"]))
    cap_data = [
        ["Tipo de captación", capture.get("tipo")],
        ["Coordenada UTM Norte", capture.get("utm_norte")],
        ["Coordenada UTM Este", capture.get("utm_este")],
        ["Datum / Huso", f"{safe_text(capture.get('datum'))} / {safe_text(capture.get('huso'))}"],
        ["Condición", capture.get("condicion")],
        ["Profundidad total", fmt(capture.get("profundidad_total"), " m")],
        ["Diámetro perforación", capture.get("diametro_perforacion")],
        ["Diámetro entubación", capture.get("diametro_entubacion")],
        ["Material / espesor tubería", capture.get("material_tuberia")],
        ["Altura sobre terreno", capture.get("altura_sobre_terreno")],
        ["Nivel estático inicial", fmt(capture.get("nivel_estatico"), " m")],
        ["Cribas / ranuras", (f"Desde {capture.get('criba_desde')} m hasta {capture.get('criba_hasta')} m" if capture.get("criba_desde") and capture.get("criba_hasta") else "No informado")],
        ["Tubería ciega", capture.get("tuberia_ciega")],
        ["Profundidad de bomba", fmt(capture.get("profundidad_bomba"), " m")],
        ["Tubería extracción/succión", capture.get("tuberia_extraccion")],
    ]
    story.append(make_table(cap_data, col_widths=[5.2 * cm, 11.2 * cm]))

    if location_image is not None:
        loc_flow = uploaded_image_flowable(location_image, width_cm=16.0, height_cm=9.0)
        if loc_flow:
            story.append(Spacer(1, 0.35 * cm))
            story.append(loc_flow)
            story.append(Paragraph("Figura 1. Croquis o imagen de ubicación de la captación.", styles["FigureCaption"]))

    story.append(Paragraph("4. Esquema constructivo", styles["SectionTitle"]))
    scheme_flow = uploaded_image_flowable(scheme_image, width_cm=9.0, height_cm=12.0) if scheme_image else None
    if scheme_flow is None:
        scheme_buf = make_simple_well_scheme(capture, stratigraphy_df)
        if scheme_buf is not None:
            scheme_flow = rl_image_preserve_aspect(scheme_buf, max_width_cm=8.2, max_height_cm=12.0)
            story.append(Paragraph("Esquema referencial generado automáticamente a partir de los datos ingresados. No reemplaza plano constructivo real.", styles["Small"]))
        else:
            scheme_flow = Paragraph("Esquema constructivo no generado: falta profundidad total o datos mínimos de captación.", styles["Body"])
    story.append(scheme_flow)
    story.append(Paragraph("Figura 2. Esquema constructivo referencial de la captación.", styles["FigureCaption"]))

    # -------------------------------------------------------------------------
    # 5. ESTRATIGRAFÍA
    # -------------------------------------------------------------------------
    story.append(Paragraph("5. Estratigrafía", styles["SectionTitle"]))
    story.append(Paragraph(
        "La estratigrafía ingresada se presenta como antecedente descriptivo del material perforado o reconocido durante la habilitación.",
        styles["Body"]
    ))
    story.append(df_to_pdf_table(stratigraphy_df, max_rows=35, font_size=6.4))

    # -------------------------------------------------------------------------
    # 6. EQUIPOS Y METODOLOGÍA
    # -------------------------------------------------------------------------
    story.append(Paragraph("6. Equipos utilizados", styles["SectionTitle"]))
    eq_data = [
        ["Bomba", equipment.get("bomba")],
        ["Potencia", equipment.get("potencia")],
        ["Tubería de extracción", capture.get("tuberia_extraccion")],
        ["Medidor de caudal", equipment.get("medidor_caudal")],
        ["Instrumento de nivel", equipment.get("instrumento_nivel")],
        ["Generador", equipment.get("generador")],
    ]
    story.append(make_table(eq_data, col_widths=[5.2 * cm, 11.2 * cm]))

    story.append(Paragraph("7. Parámetros metodológicos registrados", styles["SectionTitle"]))
    met_data = [
        ["Modo de prueba", methodology.get("modo_prueba")],
        ["Fecha de prueba", methodology.get("fecha_prueba")],
        ["Hora inicio bombeo", methodology.get("hora_inicio")],
        ["Hora término bombeo", methodology.get("hora_termino")],
        ["Duración registrada", fmt(calculations.get("duration_min"), " min", decimals=0)],
        ["Caudal objetivo", methodology.get("caudal_objetivo")],
        ["Método de medición de caudal", methodology.get("metodo_caudal")],
        ["Método de medición de niveles", methodology.get("metodo_nivel")],
        ["Frecuencia de medición", methodology.get("frecuencia")],
    ]
    story.append(make_table(met_data, col_widths=[5.2 * cm, 11.2 * cm]))

    # -------------------------------------------------------------------------
    # 8. RESULTADOS
    # -------------------------------------------------------------------------
    story.append(Paragraph("8. Resultados calculados", styles["SectionTitle"]))
    results_data = [
        ["Caudal promedio", fmt(calculations.get("q_mean"), " L/s")],
        ["Caudal mínimo", fmt(calculations.get("q_min"), " L/s")],
        ["Caudal máximo", fmt(calculations.get("q_max"), " L/s")],
        ["Variación relativa de caudal", fmt(calculations.get("q_var_pct"), " %")],
        ["Nivel estático inicial", fmt(capture.get("nivel_estatico"), " m")],
        ["Nivel dinámico final", fmt(calculations.get("final_dynamic"), " m")],
        ["Abatimiento final", fmt(calculations.get("drawdown"), " m")],
        ["Caudal específico", fmt(calculations.get("specific_capacity"), " L/s/m")],
        ["Volumen bombeado", fmt(calculations.get("volume_m3"), " m³")],
        ["Pendiente final", fmt(calculations.get("slope_cm_h"), " cm/h")],
        ["Evaluación estabilización", calculations.get("stabilization_message")],
        ["Recuperación máxima", fmt(calculations.get("recovery_max"), " %")],
        ["Tiempo a 75% recuperación", fmt(calculations.get("t75"), " min", decimals=0)],
        ["Tiempo a 90% recuperación", fmt(calculations.get("t90"), " min", decimals=0)],
        ["Tiempo a 100% recuperación", fmt(calculations.get("t100"), " min", decimals=0)],
    ]
    story.append(make_table(results_data, col_widths=[5.2 * cm, 11.2 * cm]))

    # -------------------------------------------------------------------------
    # 9. GRÁFICOS
    # -------------------------------------------------------------------------
    story.append(Paragraph("9. Gráficos", styles["SectionTitle"]))

    pump_chart = make_line_chart_image(
        pumping_df, "Tiempo_min", "Nivel_m",
        "Prueba de gasto constante",
        "Tiempo (min)", "Nivel/profundidad (m)",
        invert_y=True
    )
    story.append(image_flowable(pump_chart, width_cm=18.0, height_cm=10.0))
    story.append(Paragraph("Figura 3. Gráfico de prueba a caudal constante.", styles["FigureCaption"]))

    rec_chart = make_line_chart_image(
        recovery_df, "Tiempo_min", "Nivel_m",
        "Prueba de recuperación",
        "Tiempo (min)", "Nivel/profundidad (m)",
        invert_y=True
    )
    story.append(image_flowable(rec_chart))
    story.append(Paragraph("Figura 4. Gráfico de recuperación de nivel.", styles["FigureCaption"]))

    if recovery_df is not None and "Recuperacion_pct" in recovery_df.columns:
        rec_pct_chart = make_line_chart_image(
            recovery_df, "Tiempo_min", "Recuperacion_pct",
            "Porcentaje de recuperación",
            "Tiempo (min)", "Recuperación (%)",
            invert_y=False
        )
        story.append(image_flowable(rec_pct_chart))
        story.append(Paragraph("Figura 5. Porcentaje de recuperación acumulada.", styles["FigureCaption"]))

    # -------------------------------------------------------------------------
    # 10. TABLAS
    # -------------------------------------------------------------------------
    story.append(Paragraph("10. Tabla de prueba de gasto constante", styles["SectionTitle"]))
    story.append(df_to_pdf_table(pumping_df, max_rows=70, font_size=5.8))

    story.append(Paragraph("11. Tabla de recuperación", styles["SectionTitle"]))
    story.append(df_to_pdf_table(recovery_df, max_rows=70, font_size=5.8))

    # -------------------------------------------------------------------------
    # 12. CONCLUSIONES
    # -------------------------------------------------------------------------
    story.append(Paragraph("12. Conclusiones", styles["SectionTitle"]))
    conclusions = generate_conclusions(capture, calculations, warnings, methodology.get("modo_prueba", ""))
    for c in conclusions:
        story.append(keep_paragraph(f"• {c}", styles["Body"]))

    # Bloque de firma profesional, centrado y separado del texto final.
    story.append(Spacer(1, 2.0 * cm))

    sig_title_style = ParagraphStyle(
        name="SignatureTitle",
        parent=styles["Body"],
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=16.5,
        spaceAfter=6,
    )
    sig_text_style = ParagraphStyle(
        name="SignatureText",
        parent=styles["Body"],
        alignment=TA_CENTER,
        fontName="Helvetica",
        fontSize=11,
        leading=16.5,
        spaceAfter=0,
    )
    sig_name_style = ParagraphStyle(
        name="SignatureName",
        parent=sig_text_style,
        fontName="Helvetica-Bold",
    )

    if signature_image is not None:
        sig_flow = uploaded_image_flowable(signature_image, width_cm=5.5, height_cm=2.4)
        if sig_flow:
            sig_table_img = Table([[sig_flow]], colWidths=[8.5 * cm])
            sig_table_img.hAlign = "CENTER"
            sig_table_img.setStyle(TableStyle([
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))
            story.append(sig_table_img)

    story.append(Paragraph("____________________________", sig_text_style))
    story.append(Paragraph("David Gutiérrez Jara", sig_name_style))
    story.append(Paragraph("Ingeniero Agrónomo", sig_text_style))

    doc.build(story, onFirstPage=later_pages, onLaterPages=later_pages)
    return buffer.getvalue()


# =============================================================================
# INTERFAZ STREAMLIT
# =============================================================================

st.title("Sistema de Pruebas de Bombeo")
st.caption("Irrisal Consulting Ltda. | Informe técnico profesional v2.5.7 - firma limpia y saltos")

if LOGO_PATH.exists():
    st.image(str(LOGO_PATH), width=260)

# =============================================================================
# GUARDAR / CARGAR FICHAS DE DATOS
# =============================================================================

TIPOS_CAPTACION = ["Pozo profundo", "Noria / pozo de gran diámetro", "Puntera", "Dren", "Otro"]
CONDICIONES = ["No informado", "No surgente", "Surgente"]
MODOS_PRUEBA = ["Ensayo abreviado 180 min + recuperación", "DGA estándar 24 h", "Otro"]

def default_stratigraphy_df():
    return pd.DataFrame({
        "Desde_m": [0.0, 1.0, 5.0],
        "Hasta_m": [1.0, 5.0, 12.0],
        "Descripcion": ["Suelo vegetal", "Material fino / arcilloso", "Arena / grava / material permeable"],
        "Observacion": ["", "", ""],
    })


def default_pumping_df(mode: str):
    if "24 h" in safe_text(mode, ""):
        default_times = [0,1,2,3,4,5,10,20,30,60,120,180,240,300,360,420,480,540,600,660,720,780,840,900,960,1020,1080,1140,1200,1260,1320,1380,1440]
    else:
        default_times = [0,1,2,3,4,5,10,15,20,30,45,60,90,120,150,180]
    return pd.DataFrame({
        "Fecha": [""] * len(default_times),
        "Hora": [""] * len(default_times),
        "Tiempo_min": default_times,
        "Nivel_m": [np.nan] * len(default_times),
        "Caudal_L_s": [np.nan] * len(default_times),
        "Observacion": [""] * len(default_times),
    })


def default_recovery_df():
    default_rec_times = [0,1,2,3,4,5,10,15,20,30,45,60,90,120,150,180]
    return pd.DataFrame({
        "Fecha": [""] * len(default_rec_times),
        "Hora": [""] * len(default_rec_times),
        "Tiempo_min": default_rec_times,
        "Nivel_m": [np.nan] * len(default_rec_times),
        "Observacion": [""] * len(default_rec_times),
    })


def init_state_defaults():
    defaults = {
        "company_empresa": COMPANY_DEFAULTS["empresa"],
        "company_direccion": COMPANY_DEFAULTS["direccion"],
        "company_celular": COMPANY_DEFAULTS["celular"],
        "company_correo": COMPANY_DEFAULTS["correo"],
        "project_nombre_proyecto": "Prueba de bombeo",
        "project_identificacion": "Captación subterránea",
        "project_cliente": "",
        "project_sector": "",
        "project_comuna": "",
        "project_region": "Región del Biobío",
        "project_consultor": "",
        "project_observaciones": "",
        "capture_tipo": "Pozo profundo",
        "capture_condicion": "No informado",
        "capture_utm_norte": "",
        "capture_utm_este": "",
        "capture_datum": "SIRGAS WGS84 / WGS84",
        "capture_huso": "18S",
        "capture_profundidad_total": 0.0,
        "capture_diametro_perforacion": "",
        "capture_diametro_entubacion": "",
        "capture_material_tuberia": "",
        "capture_altura_sobre_terreno": "",
        "capture_nivel_estatico": 0.0,
        "capture_criba_desde": 0.0,
        "capture_criba_hasta": 0.0,
        "capture_tuberia_ciega": "",
        "capture_profundidad_bomba": 0.0,
        "capture_tuberia_extraccion": "",
        "capture_observaciones": "",
        "equipment_bomba": "",
        "equipment_potencia": "",
        "equipment_medidor_caudal": "",
        "equipment_instrumento_nivel": "",
        "equipment_generador": "",
        "equipment_observaciones": "",
        "methodology_modo_prueba": "Ensayo abreviado 180 min + recuperación",
        "methodology_fecha_prueba": datetime.now().strftime("%d-%m-%Y"),
        "methodology_hora_inicio": "",
        "methodology_hora_termino": "",
        "methodology_caudal_objetivo": "",
        "methodology_metodo_caudal": "Caudalímetro",
        "methodology_metodo_nivel": "Pozómetro",
        "methodology_frecuencia": "",
        "methodology_observaciones": "",
        "data_version": 0,
    }
    for key, value in defaults.items():
        st.session_state.setdefault(key, value)
    st.session_state.setdefault("strat_df_loaded", default_stratigraphy_df())
    st.session_state.setdefault("pumping_df_loaded", None)
    st.session_state.setdefault("recovery_df_loaded", None)


def coerce_state_number(key: str, default: float = 0.0):
    """
    Evita errores de Streamlit cuando una ficha guardada trae valores numéricos
    como texto, vacío, None o 'No informado'.
    """
    value = st.session_state.get(key, default)
    try:
        if value is None:
            raise ValueError
        if isinstance(value, str):
            cleaned = value.strip().replace(",", ".")
            if cleaned == "" or cleaned.lower() in ["no informado", "dato no informado", "none", "nan"]:
                raise ValueError
            value = cleaned
        value = float(value)
        if pd.isna(value):
            raise ValueError
        st.session_state[key] = value
    except Exception:
        st.session_state[key] = float(default)


def normalize_numeric_state_fields():
    """
    Normaliza todos los campos numéricos antes de construir widgets number_input.
    """
    coerce_state_number("capture_profundidad_total", 0.0)
    coerce_state_number("capture_nivel_estatico", 0.0)
    coerce_state_number("capture_criba_desde", 0.0)
    coerce_state_number("capture_criba_hasta", 0.0)
    coerce_state_number("capture_profundidad_bomba", 0.0)



def records_to_df(records, fallback_df):
    if isinstance(records, list) and records:
        return pd.DataFrame(records)
    return fallback_df


def apply_payload_to_state(payload: dict):
    sections = {
        "company": "company",
        "project": "project",
        "capture": "capture",
        "equipment": "equipment",
        "methodology": "methodology",
    }
    for section_name, prefix in sections.items():
        values = payload.get(section_name, {})
        if isinstance(values, dict):
            for field, value in values.items():
                st.session_state[f"{prefix}_{field}"] = value

    if st.session_state.get("capture_tipo") not in TIPOS_CAPTACION:
        st.session_state["capture_tipo"] = TIPOS_CAPTACION[0]
    if st.session_state.get("capture_condicion") not in CONDICIONES:
        st.session_state["capture_condicion"] = CONDICIONES[0]
    if st.session_state.get("methodology_modo_prueba") not in MODOS_PRUEBA:
        st.session_state["methodology_modo_prueba"] = MODOS_PRUEBA[0]

    normalize_numeric_state_fields()

    # Tablas editables
    st.session_state["strat_df_loaded"] = records_to_df(payload.get("stratigraphy"), default_stratigraphy_df())
    st.session_state["pumping_df_loaded"] = records_to_df(payload.get("pumping"), default_pumping_df(st.session_state.get("methodology_modo_prueba", "")))
    st.session_state["recovery_df_loaded"] = records_to_df(payload.get("recovery"), default_recovery_df())
    st.session_state["data_version"] = int(st.session_state.get("data_version", 0)) + 1


def df_to_records_for_json(df: pd.DataFrame):
    if df is None or df.empty:
        return []
    clean = df.copy()
    clean = clean.replace({np.nan: None})
    return clean.to_dict("records")


def build_payload(company, project, capture, equipment, methodology, stratigraphy_df, pumping_df, recovery_df):
    return {
        "version": "2.5",
        "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "company": company,
        "project": project,
        "capture": capture,
        "equipment": equipment,
        "methodology": methodology,
        "stratigraphy": df_to_records_for_json(stratigraphy_df),
        "pumping": df_to_records_for_json(pumping_df),
        "recovery": df_to_records_for_json(recovery_df),
    }


init_state_defaults()
normalize_numeric_state_fields()

with st.sidebar:
    st.header("Cargar ficha guardada")
    saved_json = st.file_uploader("Subir ficha de datos (.json)", type=["json"], key="saved_json_loader")
    if st.button("Aplicar datos guardados"):
        if saved_json is None:
            st.warning("Primero sube una ficha .json guardada.")
        else:
            try:
                payload = json.loads(saved_json.getvalue().decode("utf-8"))
                apply_payload_to_state(payload)
                st.success("Datos cargados. La app se actualizará con la ficha guardada.")
                st.rerun()
            except Exception as exc:
                st.error(f"No se pudo cargar la ficha: {exc}")

    st.divider()
    st.header("Configuración empresa")
    empresa = st.text_input("Empresa", key="company_empresa")
    direccion = st.text_area("Dirección", key="company_direccion")
    celular = st.text_input("Celular", key="company_celular")
    correo = st.text_input("Correo", key="company_correo")
    st.caption("Estos datos se insertan automáticamente en portada y pie de página.")

    st.header("Firma")
    signature_image = st.file_uploader(
        "Subir firma en PNG",
        type=["png"],
        key="signature_image"
    )
    st.caption("La firma se insertará al final del Word y PDF. Por seguridad, la firma no se guarda dentro de la ficha JSON.")

company = {"empresa": empresa, "direccion": direccion, "celular": celular, "correo": correo}

tabs = st.tabs([
    "1. Proyecto",
    "2. Captación",
    "3. Estratigrafía",
    "4. Equipos y metodología",
    "5. Bombeo",
    "6. Recuperación",
    "7. Resultados e informe",
])

with tabs[0]:
    st.subheader("Datos del proyecto")
    col1, col2 = st.columns(2)
    with col1:
        nombre_proyecto = st.text_input("Nombre del proyecto", key="project_nombre_proyecto")
        identificacion = st.text_input("Identificación de captación", key="project_identificacion")
        cliente = st.text_input("Cliente / beneficiario", key="project_cliente")
        sector = st.text_input("Sector / predio", key="project_sector")
    with col2:
        comuna = st.text_input("Comuna", key="project_comuna")
        region = st.text_input("Región", key="project_region")
        consultor = st.text_input("Consultor responsable", key="project_consultor")
        observaciones_proyecto = st.text_area("Observaciones generales", key="project_observaciones")

    project = {
        "nombre_proyecto": nombre_proyecto,
        "identificacion": identificacion,
        "cliente": cliente,
        "sector": sector,
        "comuna": comuna,
        "region": region,
        "consultor": consultor,
        "observaciones": observaciones_proyecto,
    }

with tabs[1]:
    st.subheader("Captación, ubicación y habilitación")
    col1, col2, col3 = st.columns(3)

    with col1:
        tipo = st.selectbox("Tipo de captación", TIPOS_CAPTACION, key="capture_tipo")
        condicion = st.selectbox("Condición", CONDICIONES, key="capture_condicion")
        utm_norte = st.text_input("UTM Norte", key="capture_utm_norte")
        utm_este = st.text_input("UTM Este", key="capture_utm_este")
        datum = st.text_input("Datum", key="capture_datum")
        huso = st.text_input("Huso", key="capture_huso")

    with col2:
        profundidad_total = st.number_input("Profundidad total (m)", min_value=0.0, step=0.1, key="capture_profundidad_total")
        diametro_perforacion = st.text_input("Diámetro perforación", key="capture_diametro_perforacion")
        diametro_entubacion = st.text_input("Diámetro entubación", key="capture_diametro_entubacion")
        material_tuberia = st.text_input("Material / espesor tubería", key="capture_material_tuberia")
        altura_sobre_terreno = st.text_input("Altura tubería sobre terreno", key="capture_altura_sobre_terreno")
        nivel_estatico = st.number_input("Nivel estático inicial (m)", min_value=-50.0, step=0.01, key="capture_nivel_estatico")

    with col3:
        st.caption("Si no tienes información de cribas o tramo filtrante, deja estos campos en 0.")
        criba_desde = st.number_input("Criba desde (m)", min_value=0.0, step=0.1, key="capture_criba_desde")
        criba_hasta = st.number_input("Criba hasta (m)", min_value=0.0, step=0.1, key="capture_criba_hasta")
        tuberia_ciega = st.text_input("Tramos tubería ciega", key="capture_tuberia_ciega")
        profundidad_bomba = st.number_input("Profundidad bomba (m)", min_value=0.0, step=0.1, key="capture_profundidad_bomba")
        tuberia_extraccion = st.text_input("Tubería extracción/succión", key="capture_tuberia_extraccion")
        observaciones_captacion = st.text_area("Observaciones de habilitación", key="capture_observaciones")

    st.write("#### Imágenes opcionales")
    location_image = st.file_uploader("Cargar croquis o imagen de ubicación", type=["jpg", "jpeg", "png"], key="location_image")
    scheme_image = st.file_uploader("Cargar esquema constructivo del pozo/captación", type=["jpg", "jpeg", "png"], key="scheme_image")

    capture = {
        "tipo": tipo,
        "condicion": condicion,
        "utm_norte": utm_norte,
        "utm_este": utm_este,
        "datum": datum,
        "huso": huso,
        "profundidad_total": profundidad_total if profundidad_total > 0 else None,
        "diametro_perforacion": diametro_perforacion,
        "diametro_entubacion": diametro_entubacion,
        "material_tuberia": material_tuberia,
        "altura_sobre_terreno": altura_sobre_terreno,
        "nivel_estatico": nivel_estatico,
        "criba_desde": criba_desde if criba_desde > 0 else "",
        "criba_hasta": criba_hasta if criba_hasta > 0 else "",
        "tuberia_ciega": tuberia_ciega,
        "profundidad_bomba": profundidad_bomba if profundidad_bomba > 0 else None,
        "tuberia_extraccion": tuberia_extraccion,
        "observaciones": observaciones_captacion,
    }

with tabs[2]:
    st.subheader("Estratigrafía")
    st.info("Ingresa los tramos reconocidos/perforados. Esta tabla se incluirá en el informe.")
    stratigraphy_df = st.data_editor(
        st.session_state.get("strat_df_loaded", default_stratigraphy_df()),
        num_rows="dynamic",
        use_container_width=True,
        key=f"strat_editor_{st.session_state.get('data_version', 0)}",
    )

with tabs[3]:
    st.subheader("Equipos utilizados y metodología")
    col1, col2 = st.columns(2)
    with col1:
        bomba = st.text_input("Bomba: tipo/marca/modelo", key="equipment_bomba")
        potencia = st.text_input("Potencia", key="equipment_potencia")
        medidor_caudal = st.text_input("Medidor de caudal", key="equipment_medidor_caudal")
        instrumento_nivel = st.text_input("Instrumento de medición de nivel", key="equipment_instrumento_nivel")
        generador = st.text_input("Generador / fuente eléctrica", key="equipment_generador")
        observaciones_equipos = st.text_area("Observaciones de equipos", key="equipment_observaciones")

    with col2:
        modo_prueba = st.selectbox("Modo de prueba", MODOS_PRUEBA, key="methodology_modo_prueba")
        fecha_prueba = st.text_input("Fecha de prueba", key="methodology_fecha_prueba")
        hora_inicio = st.text_input("Hora inicio bombeo", key="methodology_hora_inicio")
        hora_termino = st.text_input("Hora término bombeo", key="methodology_hora_termino")
        caudal_objetivo = st.text_input("Caudal objetivo", key="methodology_caudal_objetivo")
        metodo_caudal = st.text_input("Método medición caudal", key="methodology_metodo_caudal")
        metodo_nivel = st.text_input("Método medición niveles", key="methodology_metodo_nivel")
        frecuencia = st.text_input("Frecuencia de medición", key="methodology_frecuencia")
        observaciones_metodologia = st.text_area("Observaciones metodológicas", key="methodology_observaciones")

    equipment = {
        "bomba": bomba,
        "potencia": potencia,
        "medidor_caudal": medidor_caudal,
        "instrumento_nivel": instrumento_nivel,
        "generador": generador,
        "observaciones": observaciones_equipos,
    }

    methodology = {
        "modo_prueba": modo_prueba,
        "fecha_prueba": fecha_prueba,
        "hora_inicio": hora_inicio,
        "hora_termino": hora_termino,
        "caudal_objetivo": caudal_objetivo,
        "metodo_caudal": metodo_caudal,
        "metodo_nivel": metodo_nivel,
        "frecuencia": frecuencia,
        "observaciones": observaciones_metodologia,
    }

with tabs[4]:
    st.subheader("Prueba de gasto constante")
    st.warning("El sistema no rellena datos faltantes. Solo calcula con datos ingresados/importados.")
    pumping_base = st.session_state.get("pumping_df_loaded")
    if pumping_base is None:
        pumping_base = default_pumping_df(st.session_state.get("methodology_modo_prueba", ""))
    pumping_df = st.data_editor(
        pumping_base,
        num_rows="dynamic",
        use_container_width=True,
        key=f"pumping_editor_{st.session_state.get('data_version', 0)}",
    )

with tabs[5]:
    st.subheader("Prueba de recuperación")
    recovery_base = st.session_state.get("recovery_df_loaded")
    if recovery_base is None:
        recovery_base = default_recovery_df()
    recovery_df = st.data_editor(
        recovery_base,
        num_rows="dynamic",
        use_container_width=True,
        key=f"recovery_editor_{st.session_state.get('data_version', 0)}",
    )

with tabs[6]:
    st.subheader("Resultados, gráficos e informe")

    calculations, recovery_with_pct = build_calculations(pumping_df, recovery_df, nivel_estatico)

    warnings = []
    add_warning(warnings, tipo == "Pozo profundo" and calculations.get("duration_min") is not None and calculations.get("duration_min") < 1440,
                "Pozo profundo con duración menor a 24 horas: no declarar cumplimiento formal de prueba estándar de 24 h.")
    add_warning(warnings, tipo == "Puntera" and "piez" not in safe_text(instrumento_nivel, "").lower(),
                "En punteras, el control de niveles debe efectuarse en piezómetro habilitado.")
    add_warning(warnings, not utm_norte or not utm_este, "Faltan coordenadas UTM.")
    add_warning(warnings, calculations.get("stabilization_evaluable") is False, calculations.get("stabilization_message"))
    add_warning(warnings, calculations.get("is_constant") is False, "El caudal no se mantuvo constante dentro de la tolerancia definida.")
    add_warning(warnings, calculations.get("recovery_max") is not None and calculations.get("recovery_max") < 75, "Recuperación inferior a 75%; interpretación limitada.")

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Duración", fmt(calculations.get("duration_min"), " min", decimals=0))
    c2.metric("Caudal promedio", fmt(calculations.get("q_mean"), " L/s"))
    c3.metric("Abatimiento", fmt(calculations.get("drawdown"), " m"))
    c4.metric("Caudal específico", fmt(calculations.get("specific_capacity"), " L/s/m"))

    c5, c6, c7, c8 = st.columns(4)
    c5.metric("Volumen bombeado", fmt(calculations.get("volume_m3"), " m³"))
    c6.metric("Pendiente final", fmt(calculations.get("slope_cm_h"), " cm/h"))
    c7.metric("Recuperación máxima", fmt(calculations.get("recovery_max"), " %"))
    c8.metric("Tiempo 90%", fmt(calculations.get("t90"), " min", decimals=0))

    st.write("### Evaluación de estabilización")
    if calculations.get("stabilization_meets"):
        st.success(calculations.get("stabilization_message"))
    elif calculations.get("stabilization_evaluable"):
        st.error(calculations.get("stabilization_message"))
    else:
        st.warning(calculations.get("stabilization_message"))

    st.write("### Advertencias técnicas internas")
    st.caption("Estas advertencias se muestran en la app para control técnico, pero no se incluyen como sección en el Word/PDF.")
    if warnings:
        for w in warnings:
            st.warning(w)
    else:
        st.success("Sin advertencias críticas con los datos ingresados.")

    st.write("### Gráficos")
    g1, g2 = st.columns(2)

    pump_valid = df_numeric(pumping_df, ["Tiempo_min", "Nivel_m"]).dropna(subset=["Tiempo_min", "Nivel_m"])
    rec_valid = df_numeric(recovery_with_pct, ["Tiempo_min", "Nivel_m"]).dropna(subset=["Tiempo_min", "Nivel_m"])

    with g1:
        if len(pump_valid) >= 2:
            fig1 = px.line(
                pump_valid.sort_values("Tiempo_min"), x="Tiempo_min", y="Nivel_m", markers=True,
                title="Nivel/profundidad vs tiempo de bombeo",
                labels={"Tiempo_min": "Tiempo (min)", "Nivel_m": "Nivel/profundidad (m)"}
            )
            fig1.update_yaxes(autorange="reversed")
            st.plotly_chart(fig1, use_container_width=True)
        else:
            st.info("No hay suficientes datos para gráfico de bombeo.")

    with g2:
        if len(rec_valid) >= 2:
            fig2 = px.line(
                rec_valid.sort_values("Tiempo_min"), x="Tiempo_min", y="Nivel_m", markers=True,
                title="Nivel/profundidad vs tiempo de recuperación",
                labels={"Tiempo_min": "Tiempo (min)", "Nivel_m": "Nivel/profundidad (m)"}
            )
            fig2.update_yaxes(autorange="reversed")
            st.plotly_chart(fig2, use_container_width=True)
        else:
            st.info("No hay suficientes datos para gráfico de recuperación.")

    st.write("### Guardar ficha de datos")
    st.caption("Descarga esta ficha .json para recuperar la información después de actualizar la app o para reutilizar datos del mismo usuario/captación.")
    payload = build_payload(company, project, capture, equipment, methodology, stratigraphy_df, pumping_df, recovery_df)
    json_bytes = json.dumps(payload, ensure_ascii=False, indent=2, default=str).encode("utf-8")
    ficha_name = safe_text(project.get("cliente"), "ficha").lower().replace(" ", "_")
    st.download_button(
        "Guardar ficha de datos (.json)",
        data=json_bytes,
        file_name=f"ficha_prueba_bombeo_{ficha_name}.json",
        mime="application/json",
    )

    st.write("### Exportar informe")

    word_bytes = make_word_docx(
        company=company,
        project=project,
        capture=capture,
        stratigraphy_df=stratigraphy_df,
        equipment=equipment,
        methodology=methodology,
        pumping_df=pumping_df,
        recovery_df=recovery_with_pct,
        calculations=calculations,
        warnings=warnings,
        location_image=location_image,
        scheme_image=scheme_image,
        signature_image=signature_image,
    )

    col_word, col_pdf = st.columns(2)

    with col_word:
        st.download_button(
            "Descargar informe Word (.docx)",
            data=word_bytes,
            file_name="informe_prueba_bombeo.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    pdf_bytes = make_pdf(
        company=company,
        project=project,
        capture=capture,
        stratigraphy_df=stratigraphy_df,
        equipment=equipment,
        methodology=methodology,
        pumping_df=pumping_df,
        recovery_df=recovery_with_pct,
        calculations=calculations,
        warnings=warnings,
        location_image=location_image,
        scheme_image=scheme_image,
        signature_image=signature_image,
    )

    with col_pdf:
        st.download_button(
            "Descargar informe PDF (.pdf)",
            data=pdf_bytes,
            file_name="informe_prueba_bombeo.pdf",
            mime="application/pdf",
        )
